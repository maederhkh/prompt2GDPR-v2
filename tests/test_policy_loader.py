"""Standalone assert tests for the policy input loader."""
import os
import shutil
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document

from utils.policy_loader import load_policy_text, SUPPORTED_EXTENSIONS

FIXTURES = Path(__file__).parent / "fixtures"
KNOWN = "We process your personal data to provide the service."


def test_supported_extensions():
    for ext in (".txt", ".md", ".html", ".htm", ".pdf", ".docx"):
        assert ext in SUPPORTED_EXTENSIONS


def test_txt():
    d = Path(tempfile.mkdtemp())
    try:
        f = d / "p.txt"
        f.write_text(KNOWN, encoding="utf-8")
        assert KNOWN in load_policy_text(f)
    finally:
        shutil.rmtree(d)


def test_html_strips_tags_and_keeps_headings():
    d = Path(tempfile.mkdtemp())
    try:
        f = d / "p.html"
        f.write_text(
            "<html><head><style>body{color:red}</style>"
            "<script>var x=1;</script></head><body>"
            "<h2>Purpose of Processing</h2>"
            f"<p>{KNOWN}</p></body></html>",
            encoding="utf-8",
        )
        out = load_policy_text(f)
        assert KNOWN in out
        assert "Purpose of Processing" in out
        assert "<" not in out and ">" not in out      # tags gone
        assert "var x" not in out                      # script contents gone
        # heading sits on its own line (section splitter is line-based)
        assert any(line.strip() == "Purpose of Processing" for line in out.splitlines())
    finally:
        shutil.rmtree(d)


def test_pdf_fixture():
    out = load_policy_text(FIXTURES / "sample_policy.pdf")
    assert "provide the service" in out


def test_docx_generated():
    d = Path(tempfile.mkdtemp())
    try:
        f = d / "p.docx"
        doc = Document()
        doc.add_heading("Purpose of Processing", level=2)
        doc.add_paragraph(KNOWN)
        doc.save(str(f))
        out = load_policy_text(f)
        assert KNOWN in out
        assert "Purpose of Processing" in out
    finally:
        shutil.rmtree(d)


def test_unsupported_extension():
    d = Path(tempfile.mkdtemp())
    try:
        f = d / "p.rtf"
        f.write_text("hello", encoding="utf-8")
        try:
            load_policy_text(f)
            assert False, "expected ValueError for unsupported extension"
        except ValueError as exc:
            assert ".rtf" in str(exc)
    finally:
        shutil.rmtree(d)


def test_normalization():
    d = Path(tempfile.mkdtemp())
    try:
        f = d / "p.txt"
        # CRLF line endings, a non-breaking space, and a run of 4 blank lines
        raw = "line one  \r\n" + "x y\r\n" + "\n\n\n\n" + "end"
        f.write_bytes(raw.encode("utf-8"))
        out = load_policy_text(f)
        assert "\r" not in out                 # CRLF -> LF
        assert " " not in out             # nbsp removed
        assert "x y" in out                    # nbsp became a normal space
        assert "line one\n" in out             # trailing spaces stripped
        assert "\n\n\n" not in out             # 3+ blank lines collapsed
    finally:
        shutil.rmtree(d)


if __name__ == "__main__":
    test_supported_extensions()
    test_txt()
    test_html_strips_tags_and_keeps_headings()
    test_pdf_fixture()
    test_docx_generated()
    test_unsupported_extension()
    test_normalization()
    print("OK")
